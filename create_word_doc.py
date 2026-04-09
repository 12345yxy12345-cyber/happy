from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('单词格式说明 📝', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 单词数据结构
doc.add_heading('单词数据结构', level=1)
doc.add_paragraph('每个单词是一个 Python 字典，包含以下 5 个字段：')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '字段'
hdr_cells[1].text = '说明'

data = [
    ('word', '英文单词（小写）'),
    ('phonetic', '国际音标（斜杠包裹）'),
    ('pos', '词性（简写 + 点）'),
    ('meaning', '中文释义（分号分隔多个意思）'),
    ('example', '英文例句（以句号结尾）'),
]

for field, desc in data:
    row_cells = table.add_row().cells
    row_cells[0].text = field
    row_cells[1].text = desc

# 代码示例
doc.add_heading('代码格式', level=1)
code_para = doc.add_paragraph()
code_run = code_para.add_run('''
{
    "word": "单词",
    "phonetic": "/音标/",
    "pos": "词性",
    "meaning": "中文释义",
    "example": "英文例句."
}
''')
code_run.font.name = 'Courier New'
code_run.font.size = Pt(10)
code_para.paragraph_format.left_indent = Inches(0.5)

# 字段详解
doc.add_heading('字段详解', level=1)
detail_table = doc.add_table(rows=1, cols=3)
detail_table.style = 'Table Grid'
hdr_cells = detail_table.rows[0].cells
hdr_cells[0].text = '字段'
hdr_cells[1].text = '说明'
hdr_cells[2].text = '必填'

detail_data = [
    ('word', '英文单词（小写）', '是'),
    ('phonetic', '国际音标（斜杠包裹）', '是'),
    ('pos', '词性（简写 + 点）', '是'),
    ('meaning', '中文释义（分号分隔多个意思）', '是'),
    ('example', '英文例句（以句号结尾）', '是'),
]

for field, desc, required in detail_data:
    row_cells = detail_table.add_row().cells
    row_cells[0].text = field
    row_cells[1].text = desc
    row_cells[2].text = required

# 词性缩写参考
doc.add_heading('词性缩写参考', level=1)
pos_table = doc.add_table(rows=1, cols=3)
pos_table.style = 'Table Grid'
hdr_cells = pos_table.rows[0].cells
hdr_cells[0].text = '缩写'
hdr_cells[1].text = '完整形式'
hdr_cells[2].text = '中文'

pos_data = [
    ('n.', 'noun', '名词'),
    ('v.', 'verb', '动词'),
    ('adj.', 'adjective', '形容词'),
    ('adv.', 'adverb', '副词'),
    ('prep.', 'preposition', '介词'),
    ('pron.', 'pronoun', '代词'),
    ('conj.', 'conjunction', '连词'),
    ('int.', 'interjection', '感叹词'),
]

for abbr, full, cn in pos_data:
    row_cells = pos_table.add_row().cells
    row_cells[0].text = abbr
    row_cells[1].text = full
    row_cells[2].text = cn

# 完整示例
doc.add_heading('完整示例', level=1)

# 示例 1
doc.add_heading('示例 1：动词', level=2)
code_para1 = doc.add_paragraph()
code_run1 = code_para1.add_run('''
{
    "word": "abandon",
    "phonetic": "/əˈbændən/",
    "pos": "v.",
    "meaning": "放弃；抛弃；遗弃",
    "example": "He abandoned his car in the snow."
}
''')
code_run1.font.name = 'Courier New'
code_run1.font.size = Pt(10)
code_para1.paragraph_format.left_indent = Inches(0.5)

# 示例 2
doc.add_heading('示例 2：名词', level=2)
code_para2 = doc.add_paragraph()
code_run2 = code_para2.add_run('''
{
    "word": "ability",
    "phonetic": "/əˈbɪləti/",
    "pos": "n.",
    "meaning": "能力；才能；本领",
    "example": "She has the ability to learn quickly."
}
''')
code_run2.font.name = 'Courier New'
code_run2.font.size = Pt(10)
code_para2.paragraph_format.left_indent = Inches(0.5)

# 示例 3
doc.add_heading('示例 3：形容词', level=2)
code_para3 = doc.add_paragraph()
code_run3 = code_para3.add_run('''
{
    "word": "active",
    "phonetic": "/ˈæktɪv/",
    "pos": "adj.",
    "meaning": "活跃的；积极的；主动的",
    "example": "He stays active by running every day."
}
''')
code_run3.font.name = 'Courier New'
code_run3.font.size = Pt(10)
code_para3.paragraph_format.left_indent = Inches(0.5)

# 添加新单词
doc.add_heading('添加新单词', level=1)
doc.add_heading('方法一：直接编辑 vocab_bot.py', level=2)
code_para4 = doc.add_paragraph()
code_run4 = code_para4.add_run('''
word_bank = [
    # 现有单词...

    # 添加新单词
    {
        "word": "your_word",
        "phonetic": "/your-phonetic/",
        "pos": "n.",
        "meaning": "你的中文释义",
        "example": "Your example sentence."
    },
]
''')
code_run4.font.name = 'Courier New'
code_run4.font.size = Pt(10)
code_para4.paragraph_format.left_indent = Inches(0.5)

# 难度分组规则
doc.add_heading('难度分组规则', level=1)
doc.add_paragraph('简单：基础高频词', style='List Number')
doc.add_paragraph('word_bank[0:5]').paragraph_format.left_indent = Inches(0.5)
doc.add_paragraph('中等：日常常用词', style='List Number')
doc.add_paragraph('word_bank[5:10]').paragraph_format.left_indent = Inches(0.5)
doc.add_paragraph('困难：进阶词汇', style='List Number')
doc.add_paragraph('word_bank[10:]').paragraph_format.left_indent = Inches(0.5)

# 工具推荐
doc.add_heading('工具推荐', level=1)
doc.add_heading('获取音标', level=2)
tool_para = doc.add_paragraph()
tool_para.add_run('• Cambridge Dictionary: https://dictionary.cambridge.org/\n')
tool_para.add_run('• 牛津词典: https://www.oxfordlearnersdictionaries.com/\n')
tool_para.add_run('• YouGlish (发音示例): https://youglish.com/\n')

doc.add_heading('获取例句', level=2)
tool_para2 = doc.add_paragraph()
tool_para2.add_run('• Reverso Context: https://context.reverso.net/\n')
tool_para2.add_run('• Linguee: https://www.linguee.com/\n')

# 注意事项
doc.add_heading('注意事项', level=1)
note_para = doc.add_paragraph(style='List Bullet')
note_run = note_para.add_run('⚠️ 音标必须用斜杠包裹：/音标/')
note_run.font.color.rgb = RGBColor(255, 0, 0)

note_para2 = doc.add_paragraph(style='List Bullet')
note_run2 = note_para2.add_run('⚠️ 例句必须以句号结尾："sentence."')
note_run2.font.color.rgb = RGBColor(255, 0, 0)

note_para3 = doc.add_paragraph(style='List Bullet')
note_run3 = note_para3.add_run('⚠️ 多个释义用分号分隔："意思1；意思2；意思3"')
note_run3.font.color.rgb = RGBColor(255, 0, 0)

note_para4 = doc.add_paragraph(style='List Bullet')
note_run4 = note_para4.add_run('⚠️ 词性缩写后加句点："n." 而不是 "n"')
note_run4.font.color.rgb = RGBColor(255, 0, 0)

note_para5 = doc.add_paragraph(style='List Bullet')
note_run5 = note_para5.add_run('✅ 单词用小写："hello" 而不是 "Hello"')
note_run5.font.color.rgb = RGBColor(0, 128, 0)

# 批量生成模板
doc.add_heading('批量生成模板', level=1)
doc.add_paragraph('复制以下模板快速添加单词：')
code_para5 = doc.add_paragraph()
code_run5 = code_para5.add_run('''
{
    "word": "",
    "phonetic": "/",
    "pos": "",
    "meaning": "",
    "example": "."
},
''')
code_run5.font.name = 'Courier New'
code_run5.font.size = Pt(10)
code_para5.paragraph_format.left_indent = Inches(0.5)

# 保存文档
doc.save('WORD_FORMAT.docx')
print("✅ Word 文档已生成：WORD_FORMAT.docx")
